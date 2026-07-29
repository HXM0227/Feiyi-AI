from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / os.getenv("T4_WORK_GUIDE_OUTPUT", "T4模块工作讲解.docx")
BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)


def font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    tc_pr.append(element)


def set_cell_text(cell, value: str, bold=False, fill=None) -> None:
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    font(paragraph.add_run(value), 10.5, bold)


def set_table_geometry(table) -> None:
    """Apply the compact-reference token widths: 1.7in + 4.8in = 9360 DXA."""
    widths = (2448, 6912)
    tbl_pr = table._tbl.tblPr
    for tag, width in (("w:tblW", 9360), ("w:tblInd", 120)):
        element = tbl_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tbl_pr.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(text), 10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    font(p.add_run(text), 16 if level == 1 else 13, True, BLUE if level == 1 else NAVY)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    font(p.add_run(text), 10.5)


def fixed_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(4.8)
    set_cell_text(table.rows[0].cells[0], "项目", True, "E8EEF5")
    set_cell_text(table.rows[0].cells[1], "说明", True, "E8EEF5")
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, True)
        set_cell_text(cells[1], right)
    set_table_geometry(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(40)
    title.paragraph_format.space_after = Pt(8)
    font(title.add_run("T4 模块工作讲解"), 24, True, NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    font(subtitle.add_run("AI 多语种智能解说平台｜多语种理解、翻译与生成"), 12, False, RGBColor(89, 89, 89))
    fixed_table(doc, [
        ("模块定位", "P0 语言服务：将有来源的知识片段转化为中英双语讲解。"),
        ("交付版本", "T4 1.0.0；默认离线 Mock，支持千问实调。"),
        ("接口地址", "POST /v1/generate；默认监听 8104 端口。"),
        ("核心原则", "有据生成、术语一致、专名保留、事实与类比分离、全程可审计。"),
    ])

    add_heading(doc, "1. T4 在整体系统中的职责", 1)
    add_body(doc, "T4 不自行建立知识库，也不自行判断跨文化传播策略。它负责把上游已经确认的检索证据、术语约束和受众策略，转换成目标语言中的可读讲解。这样能避免模型绕过 T3 检索而凭空补充事实。")
    fixed_table(doc, [
        ("输入", "T3 的 context 引用片段；T2 的图谱约束；T5 的受众和叙事指令；用户问题与目标语种。"),
        ("处理", "语言检测、术语筛选、提示词编排、模型或 Mock 生成、引用和术语校验。"),
        ("输出", "answer、used_citation_ids，以及语言、提示词版本、术语检查、模式和降级告警。"),
        ("下游", "T0 读取答案和引用 ID；T6 可把答案送去语音合成；T9 可使用审计记录复核质量。"),
    ])

    add_heading(doc, "2. 一次导览问答如何流转", 1)
    for item in [
        "T6 将文字、语音或图片输入规范成 query，并把检测到的语言交给 T0。",
        "T0 请求 T3 返回可引用的 context；没有片段时 T0 拒绝无依据生成。",
        "T0 可叠加 T2 的知识约束与 T5 的受众策略，再调用 T4 的 /v1/generate。",
        "T4 只依据 context 组织答案，要求事实携带 [CIT-xxx]，并检查译名是否符合术语库。",
        "T0 用 T3 的原始元数据把 used_citation_ids 还原为前端可展示的来源。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 目录与关键文件", 1)
    fixed_table(doc, [
        ("t4_multilingual/", "服务入口 api.py、配置 config.py、数据契约 schemas.py、生成与审计 service.py。"),
        ("data/", "中英术语库 terminology_zh_en.json 与可复现的样例检索上下文。"),
        ("docs/interface.md", "T0 对接字段、返回字段、错误和降级行为。"),
        ("examples/", "中英请求、审计样例与回译抽检报告。"),
        ("tests/", "服务和 API 验收测试。"),
        ("runtime/audit/", "运行时审计 JSONL；被 git 忽略，避免把运行数据或敏感信息提交。"),
    ])

    add_heading(doc, "4. Mock 与千问模式", 1)
    add_body(doc, "Mock 是默认模式，完全离线且输出确定，适合演示、单元测试和没有网络的联调。qwen 模式调用千问 OpenAI 兼容 Chat Completions 接口，适合真实语言表达质量验证。两种模式的 HTTP 契约完全相同。")
    fixed_table(doc, [
        ("Mock", "T4_MODE=mock。基于首个上下文片段和全部引用 ID 组织固定格式答案。"),
        ("千问", "T4_MODE=qwen，并在 .env 中配置 DASHSCOPE_API_KEY、模型名、地址、超时和温度。"),
        ("失败兜底", "千问网络失败、空响应、未知引用或术语校验失败时，返回 generator_mode=fallback_mock 和明确 warnings。"),
        ("密钥安全", "密钥仅从本地 .env 环境变量读取；.env 已被忽略，现有密钥文件不被服务读取。"),
    ])

    add_heading(doc, "5. 术语、引用与质量控制", 1)
    add_body(doc, "术语库是 JSON 文件，每条记录包含中文名、规范英文译名、专名首次出现的中文保留规则和备注。英语讲解中，剪纸会按“Paper Cutting (剪纸)”输出；这既保留原名，也让海外受众能理解它指向的工艺。")
    for item in [
        "引用校验：答案至少包含一个 [CIT-xxx]，且每一个 ID 必须来自本次 T3 context。",
        "术语校验：从问题和证据中筛出适用术语，检查目标语言的规范译名是否出现在答案中。",
        "表达边界：提示词规定只能陈述 context 支撑的事实；帮助理解的类比必须说明它是类比。",
        "审计：记录脱敏后的问题、上下文摘录、引用 ID、提示词版本、模式、术语结果与降级情况。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. 如何启动、测试和联调", 1)
    add_body(doc, "在 PowerShell 中进入 t4，激活 .venv 后安装 requirements.txt，复制 .env.example 为 .env。使用 uvicorn t4_multilingual.api:app --host 0.0.0.0 --port 8104 启动服务。默认 Mock 不需要密钥。")
    fixed_table(doc, [
        ("运行测试", "python -m unittest discover -s tests -v"),
        ("回译抽检", "python scripts/back_translation_check.py；输出 examples/back_translation_report.json。"),
        ("T0 配置", "设置 T0_MODE=http 与 T4_BASE_URL=http://127.0.0.1:8104。"),
        ("人工复核", "查看回译报告的术语、引用和人工复核结论；重点检查文化寓意是否被过度概括。"),
    ])

    add_heading(doc, "7. 人工测试的具体流程", 1)
    add_body(doc, "人工测试先完成离线 Mock，再切换到千问实测。Mock 用来确认接口和规则是否正确；千问实测用来判断真实译文是否自然、准确且稳定。每次测试都应保留请求文件、响应 JSON 和人工结论。")
    for item in [
        "步骤 1：进入 t4，执行 .\\.venv\\Scripts\\Activate.ps1；确认虚拟环境已激活。",
        "步骤 2：执行 python -m unittest discover -s tests -v。9 项测试都应显示 OK；这验证中英生成、术语、引用、错误、降级、健康检查和审计。",
        "步骤 3：保持 T4_MODE=mock，启动服务后访问 http://127.0.0.1:8104/healthz；应返回 status=ok。",
        "步骤 4：用 examples/requests/generate-paper-cutting-en.json 向 /v1/generate 发起 POST 请求，保存返回 JSON。",
        "步骤 5：把请求中的 target_language 改为 zh-CN，重复调用，确认英文问题或英文材料能够得到中文讲解。",
        "步骤 6：将 context 改为空数组，确认返回 HTTP 422；将 target_language 改为 ja，确认同样返回 HTTP 422。",
        "步骤 7：执行 python scripts/back_translation_check.py，打开 examples/back_translation_report.json，人工比对原文、目标文本和回译。",
    ]:
        add_bullet(doc, item)
    fixed_table(doc, [
        ("目标语言", "answer 必须为所选的 en 或 zh-CN，不能混杂未解释的大段另一种语言。"),
        ("术语", "terminology_check.passed 必须为 true；例如英文首次出现剪纸时应有 Paper Cutting (剪纸)。"),
        ("来源", "used_citation_ids 只能来自请求 context；每个事实性解释应带 [CIT-xxx]。"),
        ("降级", "Mock 测试中 generator_mode=mock；千问实测中期望 generator_mode=qwen 且 fallback_used=false。"),
        ("人工结论", "检查是否超出证据补充事实、是否把类比说成史实、文化寓意是否被不当泛化。"),
    ])

    add_heading(doc, "8. 如何调用千问模型", 1)
    add_body(doc, "T4 使用千问 OpenAI 兼容的 Chat Completions 接口。服务本身不会读取项目根目录的千问 api key.txt；该文件只作为人工保存的密钥来源。使用者应手动将密钥写入 t4/.env 中的 DASHSCOPE_API_KEY。.env 已被忽略，绝不能提交到版本库、截图或复制到验收材料。")
    for item in [
        "步骤 1：在 t4 目录执行 Copy-Item .env.example .env；如果 .env 已存在，只编辑它，不要覆盖已有配置。",
        "步骤 2：在 .env 中设置 T4_MODE=qwen、DASHSCOPE_API_KEY=实际密钥、T4_QWEN_MODEL=qwen-plus。其余地址、超时和温度默认即可。",
        "步骤 3：重新启动 uvicorn 服务，然后请求 GET /readyz。返回 status=ok 说明密钥已被读取；503 说明 qwen 模式缺少密钥。",
        "步骤 4：用与 Mock 相同的 generate-paper-cutting-en.json 调用 /v1/generate。真实调用成功时 generator_mode 应为 qwen。",
        "步骤 5：如果返回 fallback_mock 或 warnings，查看告警。它表示网络、模型响应、引用或术语校验失败，系统已退回到有来源约束的 Mock 答案。",
        "步骤 6：在千问模式下再次执行回译脚本。该脚本会做一次正向生成和一次回译，便于人工确认核心事实和译名未丢失。",
    ]:
        add_bullet(doc, item)
    fixed_table(doc, [
        ("密钥位置", "现有密钥文件位于项目根目录，但 T4 只从 t4/.env 的 DASHSCOPE_API_KEY 环境变量读取。"),
        ("可配置项", "T4_QWEN_BASE_URL、T4_QWEN_MODEL、T4_QWEN_TIMEOUT_SECONDS、T4_QWEN_TEMPERATURE。"),
        ("实测记录", "保留 response 中的 generator_mode、warnings、used_citation_ids；密钥绝不写入报告或审计日志。"),
        ("网络失败", "不应把 fallback_mock 当作真实模型验收通过；应排查网络、密钥、模型名和额度后重试。"),
    ])

    add_heading(doc, "9. 常见问题", 1)
    fixed_table(doc, [
        ("返回 422", "检查 target_language 是否为 en 或 zh-CN，以及 context 是否至少含一条有 citation_id 的片段。"),
        ("千问未就绪", "qwen 模式缺少 DASHSCOPE_API_KEY 时 /readyz 返回 503；检查未提交的 .env。"),
        ("出现 fallback_mock", "查看 warnings。其含义是模型调用或质量校验失败，但答案仍受提供的资料片段和引用约束。"),
        ("术语未通过", "扩展 data/terminology_zh_en.json 或调整模型提示词，使规范译名实际出现在输出中。"),
    ])
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print("work_guide_created.docx")
